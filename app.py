import os
import json
import io

from flask import Flask, render_template, request, redirect, url_for, flash, send_file
import gspread
from oauth2client.service_account import ServiceAccountCredentials
from functools import wraps
from docx import Document

# —– Configuración de Flask —–
app = Flask(__name__, template_folder='templates', static_folder='static')
# Si prefieres, pon aquí tu SECRET_KEY en producción como variable de entorno:
app.secret_key = os.environ.get('SECRET_KEY', 'cambia_esto_por_un_valor_seguro')

# —– Autenticación básica —–
def check_auth(username, password):
    return username == os.environ.get('BASIC_USER', 'admin') \
       and password == os.environ.get('BASIC_PASS', 'password')

def authenticate():
    return ('Autorización requerida.'), 401, {
        'WWW-Authenticate': 'Basic realm="Login Required"'
    }

def requires_auth(f):
    @wraps(f)
    def decorated(*args, **kwargs):
        auth = request.authorization
        if not auth or not check_auth(auth.username, auth.password):
            return authenticate()
        return f(*args, **kwargs)
    return decorated

# —– Google Sheets setup —–
SCOPES = [
    "https://www.googleapis.com/auth/spreadsheets",
    "https://www.googleapis.com/auth/drive.readonly"
]

# Determinar ruta de credenciales
if 'GOOGLE_SHEETS_JSON' in os.environ:
    # Si se pasó el JSON en una variable de entorno, lo volcamos en disco
    creds_dict = json.loads(os.environ['GOOGLE_SHEETS_JSON'])
    CRED_FILE = '/tmp/credentials.json'
    with open(CRED_FILE, 'w', encoding='utf-8') as f:
        json.dump(creds_dict, f)
else:
    # Fallback local (solo en desarrollo)
    BASE_DIR  = os.path.dirname(os.path.abspath(__file__))
    CRED_FILE = os.path.join(BASE_DIR, 'credentials.json')

creds = ServiceAccountCredentials.from_json_keyfile_name(CRED_FILE, SCOPES)
client = gspread.authorize(creds)

# Reemplaza con tu ID y pestaña reales
SHEET_ID   = '1LDhajDpQTzi0RLw8BXLTzmA1m9yRlTX_SrxC9aKLKYg'
worksheet  = client.open_by_key(SHEET_ID).worksheet('hoja')

# —– Rutas —–
@app.route('/', methods=['GET', 'POST'])
def index():
    if request.method == 'POST':
        search_id = request.form.get('search_id', '').strip()
        if not search_id:
            flash('El campo ID no puede estar vacío.', 'error')
        else:
            try:
                cell = worksheet.find(search_id, in_column=3)
                row_idx = cell.row
                headers = worksheet.row_values(1)
                values  = worksheet.row_values(row_idx)
                record = dict(zip(headers, values))
                record['row_idx'] = row_idx
                return render_template('edit.html', record=record)
            except Exception:
                flash('ID no encontrado. Try again.', 'error')
    return render_template('index.html')

@app.route('/update', methods=['POST'])
@requires_auth
def update():
    row_idx = int(request.form.get('row_idx'))
    headers = worksheet.row_values(1)

    # Aplicar actualizaciones (permite campos vacíos)
    try:
        for col_idx, header in enumerate(headers, start=1):
            new_val = request.form.get(header, '')
            worksheet.update_cell(row_idx, col_idx, new_val)
    except Exception:
        flash('Error al guardar cambios. Intenta de nuevo.', 'error')
        record = {h: request.form.get(h, '') for h in headers}
        record['row_idx'] = row_idx
        return render_template('edit.html', record=record)

    # Exportar a Word si se solicita
    if request.form.get('export'):
        record = {h: request.form.get(h, '') for h in headers}
        doc = Document()
        doc.add_heading(f'Registro ID {record.get(headers[2])}', level=1)
        for h in headers:
            doc.add_paragraph(f'{h}: {record.get(h)}')
        bio = io.BytesIO()
        doc.save(bio)
        bio.seek(0)
        filename = f'Registro_{record.get(headers[2])}.docx'
        return send_file(
            bio,
            as_attachment=True,
            download_name=filename,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )

    flash('Registro actualizado con éxito.', 'success')
    return redirect(url_for('index'))

if __name__ == '__main__':
    app.run(debug=True)
